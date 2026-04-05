from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = REPO_ROOT / "tai_lieu"

SOURCE_TARGETS = [
    (
        DOCS_DIR / "Dac_ta_he_thong_CDSS_BHYT_20260405.md",
        DOCS_DIR / "Dac_ta_he_thong_CDSS_BHYT_20260405.docx",
        "Đặc tả hệ thống CDSS BHYT",
    ),
    (
        DOCS_DIR / "Huong_dan_su_dung_CDSS_BHYT_20260405.md",
        DOCS_DIR / "Huong_dan_su_dung_CDSS_BHYT_20260405.docx",
        "Hướng dẫn sử dụng CDSS BHYT",
    ),
]


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(11)


def add_page_number(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Trang ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def set_document_defaults(document: Document, title: str) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for style_name, size, color in [
        ("Title", 20, RGBColor(150, 24, 91)),
        ("Heading 1", 16, RGBColor(28, 54, 110)),
        ("Heading 2", 13, RGBColor(28, 54, 110)),
        ("Heading 3", 11, RGBColor(65, 65, 65)),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8)
    add_page_number(section)

    document.core_properties.title = title
    document.core_properties.author = "GitHub Copilot"
    document.core_properties.subject = title
    document.core_properties.comments = "Tài liệu xuất tự động từ Markdown nguồn trong repo"


def add_cover(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        f"Xuất tự động từ tài liệu nguồn trong repo\nNgày tạo DOCX: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    subtitle_run.font.name = "Times New Roman"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    subtitle_run.font.size = Pt(11)

    document.add_paragraph()
    document.add_section(WD_SECTION_START.NEW_PAGE)
    add_page_number(document.sections[-1])


def apply_inline_format(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(10)
        else:
            run = paragraph.add_run(part)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(11)


def add_code_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(10)


def add_markdown_content(document: Document, source_text: str) -> None:
    in_code_block = False

    for raw_line in source_text.splitlines():
        line = raw_line.rstrip()

        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            add_code_paragraph(document, line)
            continue

        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
            continue

        if stripped.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 3")
            apply_inline_format(paragraph, stripped[4:])
            continue

        if stripped.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 2")
            apply_inline_format(paragraph, stripped[3:])
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Heading 1")
            apply_inline_format(paragraph, stripped[2:])
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            apply_inline_format(paragraph, stripped[2:])
            continue

        if re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            text = re.sub(r"^\d+\.\s+", "", stripped)
            apply_inline_format(paragraph, text)
            continue

        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.space_after = Pt(6)
        apply_inline_format(paragraph, stripped)


def render_markdown_to_docx(source_path: Path, target_path: Path, title: str) -> None:
    document = Document()
    set_document_defaults(document, title)
    add_cover(document, title)
    content = source_path.read_text(encoding="utf-8")
    add_markdown_content(document, content)
    document.save(target_path)


def main() -> None:
    generated = []
    for source_path, target_path, title in SOURCE_TARGETS:
        if not source_path.exists():
            raise FileNotFoundError(f"Không tìm thấy file nguồn: {source_path}")
        render_markdown_to_docx(source_path, target_path, title)
        generated.append(target_path.name)

    print("Generated DOCX files:")
    for name in generated:
        print(f"- {name}")


if __name__ == "__main__":
    main()